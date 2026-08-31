from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "submissions" / "initial_round" / "project_creative_book_1000.md"
OUT = ROOT / "submissions" / "initial_round" / "01-作品说明文档-待填写队伍名称.docx"

FONT = "宋体"
PLACEHOLDER = "（待按官网报名信息填写）"
WORK_NAME = "考研各专业知识库问答智能体"
DIRECTION = "千问3.5模型优化创新"


def set_run_font(run, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)


def set_paragraph(paragraph, size: float = 10.5, bold: bool = False, alignment=None, first_line=True) -> None:
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if first_line else Cm(0)
    if alignment is not None:
        paragraph.alignment = alignment
    for run in paragraph.runs:
        set_run_font(run, size, bold)


def set_cell_text(cell, text: str, size: float = 9, bold: bool = False, alignment=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            if index < len(row.cells):
                row.cells[index].width = Cm(width)
                set_cell_margins(row.cells[index])


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, 9)


def configure_section(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def add_footer(section, even: bool = False) -> None:
    footer = section.even_page_footer if even else section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if even else WD_ALIGN_PARAGRAPH.RIGHT
    add_page_field(paragraph)


def add_centered_paragraph(doc, text: str, size: float, bold=False, before=0, after=0) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)


def add_cover(doc) -> None:
    add_centered_paragraph(doc, "2026中国高校计算机大赛-人工智能创意赛", 16, True, after=36)
    add_centered_paragraph(doc, "昇腾赛道项目创意书（初赛）", 22, True, after=52)

    fields = [
        ("参赛学校", PLACEHOLDER),
        ("团队名称", PLACEHOLDER),
        ("作品名称", WORK_NAME),
        ("赛题方向", DIRECTION),
        ("联系人（队长）", PLACEHOLDER),
        ("联系电话（队长）", PLACEHOLDER),
    ]
    for label, value in fields:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(2.2)
        paragraph.paragraph_format.space_after = Pt(12)
        run = paragraph.add_run(f"{label}：{value}")
        set_run_font(run, 15, True if label == "作品名称" else False)

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(90)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("中国高校计算机大赛-人工智能创意赛（昇腾赛道）组委会编制    2026.05")
    set_run_font(run, 10.5)
    doc.add_page_break()


def add_team_information(doc) -> None:
    add_centered_paragraph(doc, "2026中国高校计算机大赛-人工智能创意赛", 16, True, after=8)
    add_centered_paragraph(doc, "参赛团队信息表", 22, True, after=10)
    note = doc.add_paragraph("注：参赛团队的报名信息须与官网报名信息保持一致。以下占位符必须在提交前全部替换。")
    set_paragraph(note, 10.5, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)

    overview = doc.add_table(rows=3, cols=2)
    overview.alignment = WD_TABLE_ALIGNMENT.CENTER
    overview.style = "Table Grid"
    set_table_widths(overview, [3.2, 11.8])
    for index, (label, value) in enumerate(
        [("作品名称", WORK_NAME), ("团队名称", PLACEHOLDER), ("参赛学校", PLACEHOLDER)]
    ):
        set_cell_text(overview.cell(index, 0), label, 10.5, True)
        set_cell_shading(overview.cell(index, 0), "E7E6E6")
        set_cell_text(overview.cell(index, 1), value, 10.5, False, WD_ALIGN_PARAGRAPH.LEFT)
    set_repeat_table_header(overview.rows[0])

    heading = doc.add_paragraph("团队队员基本信息（可跨校、跨专业组队，最多3人）")
    set_paragraph(heading, 10.5, True, first_line=False)
    members = doc.add_table(rows=4, cols=9)
    members.alignment = WD_TABLE_ALIGNMENT.CENTER
    members.style = "Table Grid"
    widths = [1.4, 1.8, 1.8, 1.8, 1.1, 1.3, 1.6, 1.7, 1.4]
    set_table_widths(members, widths)
    headers = ["姓名", "学校全称", "院（系）", "专业全称", "年级", "毕业时间", "联系电话", "邮箱", "团队分工"]
    for index, header in enumerate(headers):
        set_cell_text(members.cell(0, index), header, 8, True)
        set_cell_shading(members.cell(0, index), "E7E6E6")
    set_repeat_table_header(members.rows[0])
    for row in range(1, 4):
        for column in range(9):
            set_cell_text(members.cell(row, column), "待填", 8)

    heading = doc.add_paragraph("团队指导教师信息（指导老师须与队长同校）")
    set_paragraph(heading, 10.5, True, first_line=False)
    teacher = doc.add_table(rows=2, cols=6)
    teacher.alignment = WD_TABLE_ALIGNMENT.CENTER
    teacher.style = "Table Grid"
    set_table_widths(teacher, [1.5, 2.5, 1.5, 3.4, 2.4, 3.7])
    headers = ["姓名", "院（系）全称", "职称", "研究方向", "联系电话", "联系邮箱"]
    for index, header in enumerate(headers):
        set_cell_text(teacher.cell(0, index), header, 8.5, True)
        set_cell_shading(teacher.cell(0, index), "E7E6E6")
        set_cell_text(teacher.cell(1, index), "待填", 8.5)
    set_repeat_table_header(teacher.rows[0])

    heading = doc.add_paragraph("团队成员优势描述")
    set_paragraph(heading, 10.5, True, first_line=False)
    strengths = doc.add_paragraph(
        f"{PLACEHOLDER}。建议只填写团队真实成果、项目经历、专业能力和成员分工，突出技术实现、数据整理、材料撰写及汇报答辩等能力的互补性。"
    )
    set_paragraph(strengths, 10.5)
    doc.add_page_break()


def add_originality(doc) -> None:
    add_centered_paragraph(doc, f"《{WORK_NAME}》作品原创性声明", 18, True, after=24)
    statement = (
        "郑重声明：承诺本参赛队伍报名信息真实有效；呈交的参赛作品相关资料以及所完成的作品实物等相关成果，"
        "是本团队独立进行研究工作所取得的成果，除文中已经注明引用的内容外，本作品说明文档不包含任何其他个人"
        "或集体已经发表或撰写过的作品成果，不侵犯任何第三方的知识产权或其他权利。本声明的法律结果由本参赛队承担。"
    )
    paragraph = doc.add_paragraph(statement)
    set_paragraph(paragraph, 10.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
    for line in (
        "参赛队员签名（团队全部成员）：________________________________________",
        "日期：________年______月______日",
        "指导老师审核签名：_______________________________________________",
        "日期：________年______月______日",
    ):
        paragraph = doc.add_paragraph(line)
        paragraph.paragraph_format.space_before = Pt(14)
        set_paragraph(paragraph, 10.5, first_line=False)
    note = doc.add_paragraph(
        "注：本页可打印后签名并扫描，或使用电子签名。须保证页面内容完整、字迹清晰，并与项目创意书作为同一文档提交。"
    )
    note.paragraph_format.space_before = Pt(28)
    set_paragraph(note, 10.5)
    doc.add_page_break()


def parse_markdown() -> tuple[str, list[tuple[str, list[str]]], list[list[str]]]:
    raw = SOURCE.read_text(encoding="utf-8")
    project_line = next(line for line in raw.splitlines() if line.startswith("项目名称："))
    sections: list[tuple[str, list[str]]] = []
    current_title = ""
    current_paragraphs: list[str] = []
    table_rows: list[list[str]] = []
    for line in raw.splitlines():
        stripped = line.strip()
        if stripped.startswith("## "):
            if current_title:
                sections.append((current_title, current_paragraphs))
            current_title = stripped[3:]
            current_paragraphs = []
        elif stripped.startswith("|"):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if cells and not all(set(cell) <= {"-", ":"} for cell in cells):
                table_rows.append(cells)
        elif stripped and not stripped.startswith("#") and not stripped.startswith("项目名称："):
            current_paragraphs.append(stripped)
    if current_title:
        sections.append((current_title, current_paragraphs))
    return project_line, sections, table_rows


def add_creative_book(doc) -> None:
    project_line, sections, table_rows = parse_markdown()
    add_centered_paragraph(doc, WORK_NAME, 22, True, after=6)
    add_centered_paragraph(doc, "项目创意书正文（初赛）", 16, True, after=14)
    team = doc.add_paragraph(f"团队名称：{PLACEHOLDER}")
    set_paragraph(team, 16, True, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    project = doc.add_paragraph(project_line)
    set_paragraph(project, 10.5, True, first_line=False)

    for title, paragraphs in sections:
        heading = doc.add_paragraph(title)
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
        set_paragraph(heading, 16, True, first_line=False)
        for text in paragraphs:
            paragraph = doc.add_paragraph(text)
            set_paragraph(paragraph, 10.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
        if title == "功能与进展" and table_rows:
            table = doc.add_table(rows=len(table_rows), cols=3)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = "Table Grid"
            set_table_widths(table, [3.0, 8.7, 3.3])
            for row_index, row in enumerate(table_rows):
                for column_index in range(3):
                    value = row[column_index] if column_index < len(row) else ""
                    set_cell_text(
                        table.cell(row_index, column_index),
                        value,
                        9,
                        row_index == 0,
                        WD_ALIGN_PARAGRAPH.LEFT if column_index == 1 else WD_ALIGN_PARAGRAPH.CENTER,
                    )
                    if row_index == 0:
                        set_cell_shading(table.cell(row_index, column_index), "E7E6E6")
            set_repeat_table_header(table.rows[0])

def add_demo_appendix(doc) -> None:
    heading = doc.add_paragraph("最小可运行Demo说明（选交）")
    set_paragraph(heading, 16, True, first_line=False)
    paragraphs = [
        "随文提交源码压缩包 kaoyan-knowledge-agent-initial-round.zip。压缩包内不包含任何嵌套ZIP、Qwen3.5权重、COCO数据集或虚构的NPU实测结果。",
        "Windows环境在项目根目录运行 .\\run.ps1；Linux环境运行 ./run.sh；浏览器访问 http://127.0.0.1:7860。网页支持知识问答、真题解析、院校咨询和复习规划四种模式。",
        "HTTP Skill接口为 POST /api/skill/invoke，响应包含引用来源、风险提示、模型提供方和耗时。默认显示local-template；配置真实模型服务后才会标记为外部模型连接。",
        "本地规则回归日志只用于检验固定样例上的工程行为，不作为Qwen3.5模型精度或昇腾NPU性能结论。复赛阶段将按官方流程补充真实训练、推理、精度和性能证据。",
    ]
    for text in paragraphs:
        paragraph = doc.add_paragraph(text)
        set_paragraph(paragraph, 10.5, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)


def set_document_defaults(doc) -> None:
    doc.settings.odd_and_even_pages_header_footer = True
    for section in doc.sections:
        configure_section(section)
        add_footer(section, even=False)
        add_footer(section, even=True)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def main() -> None:
    doc = Document()
    set_document_defaults(doc)
    add_cover(doc)
    add_team_information(doc)
    add_originality(doc)
    add_creative_book(doc)
    add_demo_appendix(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
